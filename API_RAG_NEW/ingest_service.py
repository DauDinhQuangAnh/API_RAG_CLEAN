from __future__ import annotations

import hashlib
import io
import logging
import os
import re
import unicodedata
import uuid
from difflib import SequenceMatcher
from typing import Any

from fastapi import HTTPException
from docx import Document
import pdfplumber

from chunking import ProtonxSemanticChunker
from API_RAG_NEW.config import (
    DEFAULT_COLLECTION_DESCRIPTION,
    EmbeddingRuntime,
    INGEST_BATCH_SIZE,
    LOCAL_EMBEDDING_PROVIDER,
)
from API_RAG_NEW.document_structure import (
    BLOCK_BULLET,
    BLOCK_PARAGRAPH,
    BLOCK_TABLE_CAPTION,
    BLOCK_UNKNOWN,
    LogicalBlock,
    build_logical_blocks,
    clean_table_title,
    stable_parent_id,
    table_contexts_from_text,
    table_to_logical_blocks,
)
from API_RAG_NEW.rag_pipeline import (
    add_records_to_collection,
    clean_collection_name,
    stable_record_id,
)
from API_RAG_NEW.schemas import IngestResponse
from API_RAG_NEW._services_shared import (
    _clean_logical_collection_name,
    _get_collection_or_404,
    _runtime_for_provider,
    logical_collection_name,
    resolve_chunking_profile,
    storage_collection_name,
    logger,
)
from API_RAG_NEW.collection_service import (
    _collection_matches_runtime_metadata,
    _embedding_collection_metadata,
    _validate_collection_embedding_metadata,
)


def ingest_file_content(
    file_name: str,
    raw_content: bytes,
    requested_collection_name: str | None,
    provider: str = LOCAL_EMBEDDING_PROVIDER,
    chunking_profile: str | None = None,
) -> IngestResponse:
    runtime = _runtime_for_provider(provider)
    effective_chunking_profile = resolve_chunking_profile(chunking_profile)
    extension = os.path.splitext(file_name)[1].casefold()
    if extension not in {".docx", ".pdf", ".txt", ".text"}:
        raise HTTPException(
            status_code=400,
            detail="Only DOCX, PDF, TXT, and TEXT files are supported.",
        )

    final_collection_name = _resolve_collection_name(file_name, requested_collection_name)
    final_storage_name = storage_collection_name(runtime.provider, final_collection_name)
    file_hash = _content_hash(raw_content)
    collection_created = False
    existing_names = {
        collection.name for collection in runtime.chroma_client.list_collections()
    }
    if final_storage_name in existing_names:
        collection = _get_collection_or_404(runtime, final_storage_name)
        _validate_collection_embedding_metadata(runtime, collection)
    else:
        collection = runtime.chroma_client.get_or_create_collection(
            name=final_storage_name,
            metadata=_embedding_collection_metadata(
                runtime,
                DEFAULT_COLLECTION_DESCRIPTION,
                logical_name=final_collection_name,
                storage_name=final_storage_name,
            ),
        )
        collection_created = True

    chunker = ProtonxSemanticChunker(model=runtime.embedding_model)
    pending_records: list[dict[str, Any]] = []
    chunk_count = 0
    warnings: list[str] = []
    requested_profile = effective_chunking_profile
    effective_profile = requested_profile
    chunk_stats = _new_chunk_stats(effective_profile)

    try:
        try:
            source_count, records, build_stats = _build_ingest_records(
                extension,
                raw_content,
                file_name,
                file_hash,
                chunker,
                requested_profile,
            )
            _merge_chunk_stats(chunk_stats, build_stats)
            if requested_profile == "hybrid":
                records = list(records)
        except Exception as exc:
            if requested_profile != "hybrid":
                raise
            warning = (
                "Hybrid chunking failed; fell back to semantic chunking: "
                f"{exc}"
            )
            warnings.append(warning)
            effective_profile = "semantic"
            chunk_stats = _new_chunk_stats(effective_profile)
            source_count, records, build_stats = _build_ingest_records(
                extension,
                raw_content,
                file_name,
                file_hash,
                chunker,
                effective_profile,
            )
            _merge_chunk_stats(chunk_stats, build_stats)

        _validate_collection_chunking_profile(collection, effective_profile)

        for record in records:
            _update_chunk_stats(chunk_stats, record)
            pending_records.append(record)
            if len(pending_records) >= INGEST_BATCH_SIZE:
                chunk_count += add_records_to_collection(
                    pending_records,
                    runtime.embedding_model,
                    collection,
                )
                pending_records.clear()

        if pending_records:
            chunk_count += add_records_to_collection(
                pending_records,
                runtime.embedding_model,
                collection,
            )

        if chunk_count == 0:
            raise HTTPException(status_code=400, detail="No valid text to chunk.")

        _set_collection_chunking_profile(collection, runtime, effective_profile)

        return IngestResponse(
            collection_name=final_collection_name,
            rows=source_count,
            chunks=chunk_count,
            warnings=warnings,
            chunking_profile=effective_profile,
            chunk_stats=_finalize_chunk_stats(chunk_stats, chunk_count),
        )
    except Exception:
        _rollback_new_empty_collection(
            runtime,
            final_storage_name,
            collection,
            collection_created=collection_created,
            chunks_added=chunk_count,
        )
        raise


def delete_document(
    collection_name: str,
    source: str,
    provider: str = LOCAL_EMBEDDING_PROVIDER,
) -> dict[str, object]:
    runtime = _runtime_for_provider(provider)
    storage_name = storage_collection_name(runtime.provider, collection_name)
    collection = _get_collection_or_404(runtime, storage_name)
    _validate_collection_embedding_metadata(runtime, collection)
    try:
        collection.delete(where={"source": source})
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to delete document: {exc}",
        ) from exc
    return {"detail": f"Document '{source}' deleted successfully.", "source": source}


def list_documents(
    collection_name: str,
    provider: str = LOCAL_EMBEDDING_PROVIDER,
) -> dict[str, list[dict[str, Any]]]:
    runtime = _runtime_for_provider(provider)
    storage_name = storage_collection_name(runtime.provider, collection_name)
    collection = _get_collection_or_404(runtime, storage_name)
    _validate_collection_embedding_metadata(runtime, collection)

    payload = collection.get(include=["metadatas"])
    metadatas = payload.get("metadatas") or []

    sources: dict[str, dict[str, Any]] = {}
    for meta in metadatas:
        if not isinstance(meta, dict):
            continue
        source = meta.get("source")
        if not source:
            continue
        if source not in sources:
            sources[source] = {
                "source": source,
                "source_type": meta.get("source_type"),
                "chunk_count": 0,
                "doc_id": meta.get("doc_id"),
            }
        sources[source]["chunk_count"] += 1

    return {"documents": list(sources.values())}


def _rollback_new_empty_collection(
    runtime: EmbeddingRuntime,
    storage_name: str,
    collection: Any,
    *,
    collection_created: bool,
    chunks_added: int,
) -> None:
    if not collection_created or chunks_added > 0:
        return

    try:
        current_count = int(collection.count())
    except Exception as exc:
        logger.warning(
            "Skipping ingest rollback for %s because collection count failed: %s",
            storage_name,
            exc,
        )
        return

    if current_count != 0:
        logger.info(
            "Skipping ingest rollback for %s because collection count is %s.",
            storage_name,
            current_count,
        )
        return

    try:
        logger.warning("Rolling back newly created empty collection %s.", storage_name)
        runtime.chroma_client.delete_collection(name=storage_name)
    except Exception as exc:
        logger.warning(
            "Failed to roll back newly created empty collection %s: %s",
            storage_name,
            exc,
        )


def _build_ingest_records(
    extension: str,
    raw_content: bytes,
    file_name: str,
    file_hash: str,
    chunker: ProtonxSemanticChunker,
    chunking_profile: str,
) -> tuple[int, Any, dict[str, Any]]:
    if extension == ".pdf":
        if chunking_profile == "hybrid":
            pages_with_tables = _extract_pdf_pages_with_tables(raw_content)
            cleanup_stats = {"skipped_flattened_table_chunks": 0}
            records = list(
                _iter_hybrid_pdf_chunk_records(
                    pages_with_tables,
                    file_name,
                    extension,
                    file_hash,
                    chunker,
                    cleanup_stats=cleanup_stats,
                )
            )
            return (1, records, cleanup_stats)
        pages = _extract_pdf_pages(raw_content)
        return (
            1,
            _iter_pdf_chunk_records(pages, file_name, extension, file_hash, chunker),
            {},
        )

    text = _extract_non_pdf_document_text(file_name, raw_content, extension)
    if chunking_profile == "hybrid":
        return (
            1,
            _iter_hybrid_document_chunk_records(
                text, file_name, extension, file_hash, chunker
            ),
            {},
        )
    return (
        1,
        _iter_document_chunk_records(text, file_name, extension, file_hash, chunker),
        {},
    )


def _validate_collection_chunking_profile(
    collection: Any,
    effective_chunking_profile: str,
) -> None:
    metadata = collection.metadata or {}
    existing_profile = metadata.get("chunking_profile")
    if existing_profile is None:
        return
    if str(existing_profile).strip().casefold() == effective_chunking_profile:
        return
    raise HTTPException(
        status_code=400,
        detail=(
            "Collection was created with a different chunking_profile. "
            "Please create a new collection or re-ingest documents."
        ),
    )


def _set_collection_chunking_profile(
    collection: Any,
    runtime: EmbeddingRuntime,
    final_profile: str,
) -> None:
    merged_metadata = dict(collection.metadata or {})
    merged_metadata.update(
        _embedding_collection_metadata(
            runtime,
            None,
            final_profile,
            logical_name=logical_collection_name(
                runtime.provider,
                collection.name,
                collection.metadata,
            ),
            storage_name=collection.name,
        )
    )
    collection.modify(metadata=merged_metadata)


def _resolve_collection_name(
    file_name: str, requested_collection_name: str | None
) -> str:
    if requested_collection_name:
        try:
            return _clean_logical_collection_name(requested_collection_name)
        except HTTPException as exc:
            if exc.detail == "Invalid collection name.":
                raise HTTPException(
                    status_code=400,
                    detail="Invalid collection_name.",
                ) from exc
            raise

    base_name = clean_collection_name(os.path.splitext(file_name)[0]) or "rag_collection"
    return f"rag_collection_{base_name}_{uuid.uuid4().hex[:6]}"


def _content_hash(raw_content: bytes) -> str:
    return hashlib.sha256(raw_content).hexdigest()


def _document_id(file_hash: str) -> str:
    return f"doc_{file_hash[:32]}"


def _extract_non_pdf_document_text(
    file_name: str, raw_content: bytes, extension: str
) -> str:
    try:
        if extension == ".docx":
            return _extract_docx_text(raw_content)
        if extension in {".txt", ".text"}:
            return _decode_text_file(raw_content)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to read {extension.lstrip('.').upper()} file: {exc}",
        ) from exc

    raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_name}")


def _extract_pdf_pages(raw_content: bytes) -> list[tuple[int, str]]:
    pages: list[tuple[int, str]] = []
    with pdfplumber.open(io.BytesIO(raw_content)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            page_text = _clean_pdf_page_text(page_text)
            if page_text:
                pages.append((page_number, page_text))
    return pages


def _extract_pdf_pages_with_tables(raw_content: bytes) -> list[dict[str, Any]]:
    pages: list[dict[str, Any]] = []
    with pdfplumber.open(io.BytesIO(raw_content)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            page_text = _clean_pdf_page_text(page_text)
            try:
                raw_tables = page.extract_tables() or []
            except Exception:
                raw_tables = []

            tables = [
                {"table_index": table_index, "rows": table}
                for table_index, table in enumerate(raw_tables, start=1)
                if table
            ]
            if page_text or tables:
                pages.append(
                    {
                        "page_number": page_number,
                        "text": page_text,
                        "tables": tables,
                    }
                )
    return pages


def _clean_pdf_page_text(text: str) -> str:
    lines: list[str] = []

    for raw_line in text.splitlines():
        line = re.sub(r"[ \t]+", " ", raw_line).strip()
        if not line:
            continue
        if re.fullmatch(r"Trang\s+\d+\s*/\s*\d+", line, flags=re.IGNORECASE):
            continue
        if re.fullmatch(r"[•●○▪\-–—]+", line):
            continue
        lines.append(line)

    paragraphs: list[str] = []
    current = ""
    for line in lines:
        if not current:
            current = line
            continue

        if _should_join_pdf_line(current, line):
            current = f"{current} {line}"
        else:
            paragraphs.append(current)
            current = line

    if current:
        paragraphs.append(current)

    cleaned = "\n".join(paragraphs)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _should_join_pdf_line(previous: str, current: str) -> bool:
    if re.fullmatch(r"\d+\.", previous):
        return True
    if previous.endswith(("/", "-", "–", "—")):
        return True
    if re.search(r"[.!?;:]$", previous):
        return False
    if re.match(r"^\d+\.|^[a-zA-Z]\)", current):
        return False
    return True


def _extract_docx_text(raw_content: bytes) -> str:
    document = Document(io.BytesIO(raw_content))
    blocks: list[str] = []

    blocks.extend(paragraph.text.strip() for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(block for block in blocks if block)


def _decode_text_file(raw_content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-16", "cp1258", "latin-1"):
        try:
            return raw_content.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise HTTPException(status_code=400, detail="Failed to decode text file.")


def _iter_document_chunk_records(
    text: str,
    file_name: str,
    extension: str,
    file_hash: str,
    chunker: ProtonxSemanticChunker,
):
    if not text.strip():
        return

    doc_id = _document_id(file_hash)
    source_type = extension.lstrip(".")
    for chunk_index, chunk in enumerate(chunker.split_text(text), start=1):
        if not chunk.strip():
            continue
        yield {
            "id": stable_record_id(doc_id, file_name, source_type, chunk_index, chunk),
            "chunk": chunk,
            "doc_id": doc_id,
            "source": file_name,
            "source_type": source_type,
            "chunk_index": chunk_index,
        }


def _iter_pdf_chunk_records(
    pages: list[tuple[int, str]],
    file_name: str,
    extension: str,
    file_hash: str,
    chunker: ProtonxSemanticChunker,
):
    doc_id = _document_id(file_hash)
    source_type = extension.lstrip(".")
    chunk_index = 0
    for page_number, page_text in pages:
        page_chunk_index = 0
        for chunk in chunker.split_text(page_text):
            if not chunk.strip():
                continue
            chunk_index += 1
            page_chunk_index += 1
            yield {
                "id": stable_record_id(
                    doc_id,
                    file_name,
                    source_type,
                    page_number,
                    page_chunk_index,
                    chunk,
                ),
                "chunk": chunk,
                "doc_id": doc_id,
                "source": file_name,
                "source_type": source_type,
                "chunk_index": chunk_index,
                "page_number": page_number,
                "page_chunk_index": page_chunk_index,
            }


def _iter_hybrid_document_chunk_records(
    text: str,
    file_name: str,
    extension: str,
    file_hash: str,
    chunker: ProtonxSemanticChunker,
):
    if not text.strip():
        return

    doc_id = _document_id(file_hash)
    source_type = extension.lstrip(".")
    chunk_index = 0
    for base_record in _build_hybrid_text_chunk_records(
        text, file_name, source_type, doc_id, chunker
    ):
        chunk = base_record["chunk"]
        chunk_index += 1
        yield {
            **base_record,
            "id": stable_record_id(
                doc_id,
                file_name,
                source_type,
                base_record.get("section_path") or "",
                chunk_index,
                chunk,
            ),
            "chunk_index": chunk_index,
        }


def _iter_hybrid_pdf_chunk_records(
    pages: list[dict[str, Any]],
    file_name: str,
    extension: str,
    file_hash: str,
    chunker: ProtonxSemanticChunker,
    cleanup_stats: dict[str, Any] | None = None,
):
    doc_id = _document_id(file_hash)
    source_type = extension.lstrip(".")
    chunk_index = 0

    for page in pages:
        page_number = int(page.get("page_number") or 0)
        page_text = str(page.get("text") or "")
        tables = page.get("tables") or []
        page_chunk_index = 0

        contexts = table_contexts_from_text(page_text, page_number=page_number)
        table_base_records: list[dict[str, Any]] = []
        next_block_index = _next_table_block_index(contexts)
        for fallback_table_index, table in enumerate(tables, start=1):
            table_index = int(table.get("table_index") or fallback_table_index)
            context = _table_context_for_index(contexts, table_index)
            table_title = clean_table_title(
                context.table_title if context else None
            ) or "N/A"
            table_blocks = table_to_logical_blocks(
                table.get("rows") or [],
                table_index=table_index,
                table_title=table_title,
                page_number=page_number,
                section_title=context.section_title if context else None,
                section_path=context.section_path if context else None,
                start_block_index=next_block_index,
            )
            next_block_index += len(table_blocks)
            for block in table_blocks:
                parent_id = stable_parent_id(
                    doc_id,
                    file_name,
                    page_number,
                    block.section_path,
                    0 if block.section_path else block.block_index,
                    block.section_path or block.table_title or block.text,
                )
                record = {
                    "id": stable_record_id(
                        doc_id,
                        file_name,
                        source_type,
                        page_number,
                        block.table_index,
                        block.table_row_index,
                        block.table_row_part_index or "",
                        block.text,
                    ),
                    "chunk": block.text,
                    "doc_id": doc_id,
                    "source": file_name,
                    "source_type": source_type,
                    "page_number": page_number,
                    "chunk_type": "table_row",
                    "section_title": block.section_title,
                    "section_path": block.section_path,
                    "block_index": block.block_index,
                    "parent_id": parent_id,
                    "table_index": block.table_index,
                    "table_title": block.table_title or table_title,
                    "table_row_index": block.table_row_index,
                }
                if block.table_row_part_index is not None:
                    record["table_row_part_index"] = block.table_row_part_index
                table_base_records.append(record)

        semantic_candidates = _build_hybrid_text_chunk_records(
            page_text,
            file_name,
            source_type,
            doc_id,
            chunker,
            page_number=page_number,
            include_table_captions=not bool(tables),
        )
        semantic_records, skipped_count = _filter_flattened_table_semantic_records(
            semantic_candidates,
            table_base_records,
        )
        if cleanup_stats is not None:
            cleanup_stats["skipped_flattened_table_chunks"] = int(
                cleanup_stats.get("skipped_flattened_table_chunks", 0)
            ) + skipped_count

        for base_record in semantic_records:
            chunk = base_record["chunk"]
            chunk_index += 1
            page_chunk_index += 1
            yield {
                **base_record,
                "id": stable_record_id(
                    doc_id,
                    file_name,
                    source_type,
                    page_number,
                    base_record.get("section_path") or "",
                    page_chunk_index,
                    chunk,
                ),
                "chunk_index": chunk_index,
                "page_chunk_index": page_chunk_index,
            }

        for base_record in table_base_records:
            chunk_index += 1
            page_chunk_index += 1
            yield {
                **base_record,
                "chunk_index": chunk_index,
                "page_chunk_index": page_chunk_index,
            }


def _build_hybrid_text_chunk_records(
    text: str,
    file_name: str,
    source_type: str,
    doc_id: str,
    chunker: ProtonxSemanticChunker,
    *,
    page_number: int | None = None,
    include_table_captions: bool = True,
) -> list[dict[str, Any]]:
    blocks = build_logical_blocks(text, page_number=page_number)
    groups = _group_text_blocks(blocks, include_table_captions=include_table_captions)
    records: list[dict[str, Any]] = []

    for group in groups:
        group_text = "\n".join(block.text for block in group).strip()
        if not group_text:
            continue

        first_block = group[0]
        parent_id = stable_parent_id(
            doc_id,
            file_name,
            page_number,
            first_block.section_path,
            0 if first_block.section_path else first_block.block_index,
            first_block.section_path or group_text,
        )
        chunk_type = "section_child" if first_block.section_path else "fallback_semantic"
        for chunk in chunker.split_text(group_text):
            chunk = chunk.strip()
            if not chunk:
                continue
            record = {
                "chunk": chunk,
                "doc_id": doc_id,
                "source": file_name,
                "source_type": source_type,
                "chunk_type": chunk_type,
                "section_title": first_block.section_title,
                "section_path": first_block.section_path,
                "block_index": first_block.block_index,
                "parent_id": parent_id,
            }
            if page_number is not None:
                record["page_number"] = page_number
            records.append(record)

    if records or not text.strip() or blocks:
        return records

    parent_id = stable_parent_id(doc_id, file_name, page_number, None, 1, text)
    for chunk in chunker.split_text(text):
        chunk = chunk.strip()
        if not chunk:
            continue
        record = {
            "chunk": chunk,
            "doc_id": doc_id,
            "source": file_name,
            "source_type": source_type,
            "chunk_type": "fallback_semantic",
            "block_index": 1,
            "parent_id": parent_id,
        }
        if page_number is not None:
            record["page_number"] = page_number
        records.append(record)
    return records


def _filter_flattened_table_semantic_records(
    semantic_records: list[dict[str, Any]],
    table_records: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], int]:
    if not semantic_records or not table_records:
        return semantic_records, 0

    table_signatures = [
        _table_signature(str(record.get("chunk") or ""))
        for record in table_records
    ]
    table_signatures = [sig for sig in table_signatures if sig[1]]
    if not table_signatures:
        return semantic_records, 0

    kept_records: list[dict[str, Any]] = []
    skipped_count = 0
    for record in semantic_records:
        chunk = str(record.get("chunk") or "")
        if _is_flattened_table_chunk(chunk, table_signatures):
            skipped_count += 1
            continue
        kept_records.append(record)

    return kept_records, skipped_count


def _table_signature(text: str) -> tuple[str, set[str]]:
    normalized = _normalize_for_table_overlap(text)
    return normalized, _table_overlap_tokens(normalized)


def _is_flattened_table_chunk(
    chunk: str,
    table_signatures: list[tuple[str, set[str]]],
) -> bool:
    candidate_normalized, candidate_tokens = _table_signature(chunk)
    if len(candidate_tokens) < 8 or len(candidate_normalized) < 50:
        return False

    for table_normalized, table_tokens in table_signatures:
        if len(table_tokens) < 8:
            continue
        shared_tokens = candidate_tokens & table_tokens
        if len(shared_tokens) < 8:
            continue

        table_overlap = len(shared_tokens) / max(len(table_tokens), 1)
        candidate_overlap = len(shared_tokens) / max(len(candidate_tokens), 1)
        if (
            table_normalized
            and table_normalized in candidate_normalized
            and table_overlap >= 0.7
        ):
            return True
        if (
            candidate_normalized in table_normalized
            and len(candidate_normalized) >= 120
            and candidate_overlap >= 0.7
        ):
            return True
        if table_overlap >= 0.78 and candidate_overlap >= 0.58:
            return True

        ratio = SequenceMatcher(
            None,
            candidate_normalized[:4000],
            table_normalized[:4000],
        ).ratio()
        if ratio >= 0.82 and table_overlap >= 0.62 and candidate_overlap >= 0.5:
            return True

    return False


def _normalize_for_table_overlap(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text or "")
    without_marks = "".join(
        char for char in normalized if not unicodedata.combining(char)
    )
    without_marks = without_marks.casefold()
    without_marks = re.sub(r"[^a-z0-9]+", " ", without_marks)
    return re.sub(r"\s+", " ", without_marks).strip()


def _table_overlap_tokens(normalized_text: str) -> set[str]:
    stop_tokens = {"table", "row", "column", "bang", "hang", "cot", "n/a", "na"}
    return {
        token
        for token in normalized_text.split()
        if len(token) >= 3 and token not in stop_tokens
    }


def _group_text_blocks(
    blocks: list[LogicalBlock],
    *,
    include_table_captions: bool,
) -> list[list[LogicalBlock]]:
    groups: list[list[LogicalBlock]] = []
    current_group: list[LogicalBlock] = []
    current_key: tuple[str | None, str | None] | None = None
    text_block_types = {BLOCK_PARAGRAPH, BLOCK_BULLET, BLOCK_UNKNOWN}
    if include_table_captions:
        text_block_types.add(BLOCK_TABLE_CAPTION)

    for block in blocks:
        if block.block_type not in text_block_types:
            continue
        key = (block.section_path, block.section_title)
        if current_group and key == current_key:
            current_group.append(block)
            continue
        if current_group:
            groups.append(current_group)
        current_group = [block]
        current_key = key

    if current_group:
        groups.append(current_group)
    return groups


def _next_table_block_index(contexts: list[Any]) -> int:
    indexes = [int(context.block_index) for context in contexts if context.block_index]
    return (max(indexes) + 1) if indexes else 1


def _table_context_for_index(contexts: list[Any], table_index: int) -> Any | None:
    if not contexts:
        return None
    if 1 <= table_index <= len(contexts):
        return contexts[table_index - 1]
    return contexts[-1]


def _new_chunk_stats(profile: str) -> dict[str, Any]:
    return {
        "total_chunks": 0,
        "semantic_chunks": 0,
        "table_chunks": 0,
        "avg_chunk_chars": 0,
        "small_chunks": 0,
        "large_chunks": 0,
        "pages": 0,
        "profile": profile,
        "skipped_flattened_table_chunks": 0,
        "_total_chars": 0,
        "_pages": set(),
    }


def _merge_chunk_stats(stats: dict[str, Any], extra_stats: dict[str, Any]) -> None:
    for key, value in (extra_stats or {}).items():
        if isinstance(value, int):
            stats[key] = int(stats.get(key, 0)) + value
        elif value is not None:
            stats[key] = value


def _update_chunk_stats(stats: dict[str, Any], record: dict[str, Any]) -> None:
    chunk = str(record.get("chunk") or "")
    stats["total_chunks"] += 1
    stats["_total_chars"] += len(chunk)
    if len(chunk) < 120:
        stats["small_chunks"] += 1
    if len(chunk) > 1400:
        stats["large_chunks"] += 1

    if record.get("chunk_type") == "table_row":
        stats["table_chunks"] += 1
    else:
        stats["semantic_chunks"] += 1

    page_number = record.get("page_number")
    if page_number is not None:
        stats["_pages"].add(page_number)


def _finalize_chunk_stats(stats: dict[str, Any], chunk_count: int) -> dict[str, Any]:
    total_chars = stats.pop("_total_chars", 0)
    pages = stats.pop("_pages", set())
    stats["total_chunks"] = chunk_count
    stats["pages"] = len(pages)
    stats["avg_chunk_chars"] = round(total_chars / chunk_count, 2) if chunk_count else 0
    return stats
